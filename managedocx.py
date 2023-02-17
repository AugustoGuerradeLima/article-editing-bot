#document management
import docx
import json
#open the document .docx
doc = docx.Document('doc.docx')

#start the variables
#title<-first paragraph
title = doc.paragraphs[0].text

#info<-between the first paragraph and the word "Resumo"
info = ''
for para in doc.paragraphs[1:]:
    if 'Resumo' in para.text:
        break
    else:
        info += para.text

#abstract<-between the words Resumo and Keywords
abstract = ''
found_resumo = False
for para in doc.paragraphs:
    if 'Resumo' in para.text:
        found_resumo = True
    elif 'Keywords' in para.text:
        break
    elif found_resumo:
        abstract += para.text

#create a JSON with the variables
text = {
    "title": title,
    "author":"author"
    "info": info,
    "abstract": abstract,
    "key1":"key1"
    "entitle":"entitle"
    "key2":"key2",
    "content": "content"
}

#test
text_json = json.dumps(text, indent=4)
print(text_json)
